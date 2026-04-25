"""
extractors/word.py — python-docx Word document extraction.
"""
import re
from pathlib import Path

from docx import Document

from src.config_loader import general_settings
from src.logger import get_logger
from src.extractors.pdf import _chunk_text

logger   = get_logger()
CHUNK_SZ = general_settings["ingestion"]["chunk_size"]
OVERLAP  = general_settings["ingestion"]["chunk_overlap"]


def extract_word(file_path: str) -> dict:
    path = Path(file_path)
    doc  = Document(str(path))

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    full_text  = "\n\n".join(paragraphs)

    # Try to extract title from first heading or filename
    title = None
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading") and para.text.strip():
            title = para.text.strip()
            break
    if not title:
        title = path.stem.replace("_", " ").replace("-", " ").title()

    raw_chunks  = _chunk_text(full_text, CHUNK_SZ, OVERLAP)
    all_chunks  = [
        {
            "content":         c,
            "chunk_index":     i,
            "page_number":     None,
            "timestamp_start": None,
            "timestamp_end":   None,
            "metadata":        {},
        }
        for i, c in enumerate(raw_chunks)
    ]

    logger.info(f"WORD | {path.name} | {len(all_chunks)} chunks")
    return {
        "text":       full_text,
        "title":      title,
        "chunks":     all_chunks,
        "word_count": len(full_text.split()),
        "metadata":   {"extractor": "python-docx"},
    }